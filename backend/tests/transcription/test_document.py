import io
from datetime import datetime

from docx import Document
from schemas.transcription.minutes import Agreement, AgreementResult, MeetingType, MinutesResponse, Task
from services.transcription.document_service import DocumentService


def test_docx_generation_integrity():
    mock_data = MinutesResponse(
        title="Junta ordinaria marzo 2026",
        scheduled_at=datetime(2026, 3, 24, 19, 0),
        location="Residencial Vecinus",
        meeting_type=MeetingType.ORDINARY,
        transcription=(
            "Primer bloque de transcripcion con suficiente longitud para forzar que el documento genere varios "
            "parrafos legibles y no deje todo el contenido en un unico bloque. "
            "Se comenta la aprobacion del presupuesto y las dudas de varios vecinos.\n\n"
            "Segundo bloque de transcripcion con mas detalles sobre las tareas acordadas y el calendario de "
            "ejecucion para las siguientes semanas."
        ),
        summary="Resumen de prueba",
        topics=["Tema 1"],
        agreements=[Agreement(description="Se aprueba el presupuesto", result=AgreementResult.APPROVED)],
        tasks=[Task(responsible="Yo", description="Test", deadline="Pronto")],
    )

    buffer = DocumentService.generate_docx(mock_data)

    if not isinstance(buffer, io.BytesIO):
        raise AssertionError(
            "Tipo de salida incorrecto." + f"Se esperaba io.BytesIO pero se obtuvo {type(buffer).__name__}"
        )

    doc = Document(buffer)
    texts = [p.text for p in doc.paragraphs if p.text.strip()]

    if texts[0] != "Junta ordinaria marzo 2026":
        raise AssertionError(f"El titulo principal del documento no coincide con el titulo del acta: {texts[0]}")

    if "Acta de Reunion" in texts or "Acta de Reunión" in texts:
        raise AssertionError("El documento no debe usar un titulo generico si se ha indicado un titulo de reunion.")

    if not any("Fecha de generacion:" in text or "Fecha de generación:" in text for text in texts):
        raise AssertionError("No se encontro el subtitulo con fecha, hora y ubicacion.")

    if "Se aprueba el presupuesto Resultado: Aprobado" not in texts:
        raise AssertionError("El resultado del acuerdo no se tradujo al espanol en el DOCX.")

    if any("APPROVED" in text or "DENIED" in text for text in texts):
        raise AssertionError("El documento no debe mostrar los valores internos del enum.")

    if "Hoja de asistentes y firmas" not in texts:
        raise AssertionError("No se encontro la hoja final reservada para asistentes y firmas.")

    transcription_index = texts.index("Transcripcion Completa")
    signature_index = texts.index("Hoja de asistentes y firmas")
    transcription_paragraphs = texts[transcription_index + 1 : signature_index]
    if len(transcription_paragraphs) < 2:
        raise AssertionError("La transcripcion debe generarse en varios parrafos y no en un unico bloque.")

    if len(doc.tables) < 2:
        raise AssertionError("El documento deberia contener la tabla de tareas y la hoja de firmas.")

    if len(doc.tables[0].rows) != 2:
        raise AssertionError("La tabla de tareas deberia tener cabecera y una fila de contenido.")

    signature_headers = [cell.text for cell in doc.tables[1].rows[0].cells]
    if signature_headers != ["Nombre completo", "Cargo en la comunidad", "Firma"]:
        raise AssertionError(f"Cabecera inesperada en la hoja de firmas: {signature_headers}")

    if len(doc.tables[1].rows) != DocumentService.SIGNATURE_BLANK_ROWS + 1:
        raise AssertionError("La tabla de firmas debe tener suficientes filas para ocupar las dos paginas finales.")

    first_signature_row = doc.tables[1].rows[1]
    if first_signature_row.height is None or first_signature_row.height.cm < 1.7:
        raise AssertionError("Las filas de firma deben tener altura suficiente para dejar espacio a la firma.")


def test_build_docx_filename_uses_meeting_title_or_default():
    assert DocumentService.build_docx_filename("Junta / Marzo 2026") == "Junta Marzo 2026.docx"
    assert DocumentService.build_docx_filename("") == "acta_reunion.docx"
