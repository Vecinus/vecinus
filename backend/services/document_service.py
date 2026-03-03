import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from schemas.transcription.minutes import MinutesResponse


def _add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def _add_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.size = Pt(11)
    return paragraph


def generate_docx(minutes: MinutesResponse) -> io.BytesIO:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Acta de Reunión", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_paragraph = doc.add_paragraph()
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_paragraph.add_run(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")

    _add_heading(doc, "Transcripción Completa", level=1)
    _add_paragraph(doc, minutes.transcription)

    _add_heading(doc, "Resumen", level=1)
    _add_paragraph(doc, minutes.summary)

    _add_heading(doc, "Temas Tratados", level=1)
    for topic in minutes.topics:
        doc.add_paragraph(topic, style="List Bullet")

    _add_heading(doc, "Acuerdos Alcanzados", level=1)
    for agreement in minutes.agreements:
        doc.add_paragraph(agreement, style="List Bullet")

    _add_heading(doc, "Tareas Asignadas", level=1)

    if minutes.tasks:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        headers = ["Responsable", "Descripción", "Plazo"]
        for i, header_text in enumerate(headers):
            header_cells[i].text = header_text
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            from docx.oxml.ns import qn

            shading = header_cells[i]._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(
                qn("w:shd"),
                {
                    qn("w:fill"): "1A1A2E",
                    qn("w:val"): "clear",
                },
            )
            shading.append(shading_elm)

        for task in minutes.tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = task.responsible
            row_cells[1].text = task.description
            row_cells[2].text = task.deadline
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
