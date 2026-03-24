import io
import re

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from schemas.transcription.minutes import MinutesResponse


class DocumentService:
    DEFAULT_FILENAME_STEM = "acta_reunion"
    DEFAULT_DOCUMENT_TITLE = "Acta de reunion"
    MAX_TRANSCRIPTION_PARAGRAPH_CHARS = 280
    SIGNATURE_ROW_HEIGHT_CM = 1.8
    SIGNATURE_BLANK_ROWS = 20
    AGREEMENT_RESULT_LABELS = {
        "APPROVED": "Aprobado",
        "DENIED": "Denegado",
    }

    @staticmethod
    def _apply_style(obj, size_pt=11, bold=False, color=None):
        if not obj:
            return
        runs = obj.runs if hasattr(obj, "runs") and obj.runs else [obj.add_run()]
        for run in runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size_pt)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color

    @staticmethod
    def add_heading(doc: Document, text: str, level: int = 1):
        text_str = str(text) if text is not None else ""
        heading = doc.add_heading(text_str, level=level)
        DocumentService._apply_style(heading, size_pt=14 if level == 1 else 12, color=RGBColor(0x1A, 0x1A, 0x2E))
        return heading

    @staticmethod
    def add_styled_paragraph(doc: Document, text: str, style: str = None):
        text_str = str(text) if text is not None else ""
        paragraph = doc.add_paragraph(text_str, style=style)
        DocumentService._apply_style(paragraph, size_pt=11)
        return paragraph

    @staticmethod
    def build_docx_filename(title: str | None) -> str:
        sanitized_title = re.sub(r'[<>:"/\\|?*\x00-\x1F]', " ", (title or "").strip())
        sanitized_title = re.sub(r"\s+", " ", sanitized_title).strip().rstrip(".")
        filename_stem = sanitized_title or DocumentService.DEFAULT_FILENAME_STEM
        return f"{filename_stem}.docx"

    @staticmethod
    def _chunk_words(text: str, max_chars: int) -> list[str]:
        words = text.split()
        if not words:
            return []

        chunks = []
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if len(candidate) <= max_chars:
                current = candidate
            else:
                chunks.append(current)
                current = word
        chunks.append(current)
        return chunks

    @staticmethod
    def _chunk_text(text: str, max_chars: int) -> list[str]:
        normalized = re.sub(r"\s+", " ", text).strip()
        if not normalized:
            return []

        sentences = [sentence.strip() for sentence in re.split(r"(?<=[.!?])\s+", normalized) if sentence.strip()]
        if not sentences:
            return DocumentService._chunk_words(normalized, max_chars)

        paragraphs = []
        current = ""
        for sentence in sentences:
            candidate = f"{current} {sentence}".strip()
            if current and len(candidate) > max_chars:
                paragraphs.append(current)
                if len(sentence) > max_chars:
                    paragraphs.extend(DocumentService._chunk_words(sentence, max_chars))
                    current = ""
                else:
                    current = sentence
            else:
                current = candidate

        if current:
            paragraphs.append(current)

        return paragraphs or DocumentService._chunk_words(normalized, max_chars)

    @staticmethod
    def _format_transcription_paragraphs(transcription: str) -> list[str]:
        normalized = re.sub(r"\r\n?", "\n", transcription or "").strip()
        if not normalized:
            return []

        raw_blocks = [block.strip() for block in re.split(r"\n\s*\n+", normalized) if block.strip()]
        paragraphs = []
        for block in raw_blocks:
            paragraphs.extend(DocumentService._chunk_text(block, DocumentService.MAX_TRANSCRIPTION_PARAGRAPH_CHARS))
        return paragraphs

    @staticmethod
    def _build_signature_sheet(doc: Document):
        doc.add_page_break()
        heading = DocumentService.add_heading(doc, "Hoja de asistentes y firmas", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.space_after = Pt(6)

        intro = DocumentService.add_styled_paragraph(
            doc,
            "Reservado para que las personas asistentes anoten su nombre completo, cargo en la comunidad y firma.",
        )
        intro.paragraph_format.space_after = Pt(10)

        table = doc.add_table(rows=DocumentService.SIGNATURE_BLANK_ROWS + 1, cols=3)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        headers = ["Nombre completo", "Cargo en la comunidad", "Firma"]
        for index, header_text in enumerate(headers):
            header_cells[index].text = header_text
            for paragraph in header_cells[index].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

        for row in table.rows[1:]:
            row.height = Cm(DocumentService.SIGNATURE_ROW_HEIGHT_CM)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for cell in row.cells:
                cell.text = ""

        return table

    @staticmethod
    def generate_docx(minutes: MinutesResponse) -> io.BytesIO:
        doc = Document()

        title = doc.add_heading(minutes.title or DocumentService.DEFAULT_DOCUMENT_TITLE, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        f_date = minutes.scheduled_at.strftime("%d/%m/%Y %H:%M")
        metadata_parag = doc.add_paragraph()
        metadata_parag.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = metadata_parag.add_run(f"Fecha de generacion: {f_date} | Ubicacion: {minutes.location}")
        run.font.size = Pt(10)
        run.font.italic = True

        DocumentService.add_heading(doc, "Resumen", level=1)
        DocumentService.add_styled_paragraph(doc, minutes.summary)

        DocumentService.add_heading(doc, "Temas Tratados", level=1)
        for topic in minutes.topics:
            DocumentService.add_styled_paragraph(doc, topic, style="List Bullet")

        DocumentService.add_heading(doc, "Acuerdos Alcanzados", level=1)
        for agreement in minutes.agreements:
            result_label = DocumentService.AGREEMENT_RESULT_LABELS.get(agreement.result.value, agreement.result.value)
            agreement_text = f"{agreement.description} Resultado: {result_label}"
            DocumentService.add_styled_paragraph(doc, agreement_text, style="List Bullet")

        DocumentService.add_heading(doc, "Tareas Asignadas", level=1)
        if minutes.tasks:
            table = doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"

            header_cells = table.rows[0].cells
            headers = ["Responsable", "Descripcion", "Plazo"]
            for i, header_text in enumerate(headers):
                header_cells[i].text = header_text
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                from docx.oxml.ns import qn

                shading = header_cells[i]._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(
                    qn("w:shd"),
                    {
                        qn("w:fill"): "1A1A2E",
                        qn("w:val"): "clear",
                    },
                )
                shading.append(shading_elm)

            for task in minutes.tasks:
                row_cells = table.add_row().cells
                row_cells[0].text = task.responsible
                row_cells[1].text = task.description
                row_cells[2].text = task.deadline
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

        DocumentService.add_heading(doc, "Transcripcion Completa", level=1)
        for paragraph_text in DocumentService._format_transcription_paragraphs(minutes.transcription):
            paragraph = DocumentService.add_styled_paragraph(doc, paragraph_text)
            paragraph.paragraph_format.space_after = Pt(10)

        DocumentService._build_signature_sheet(doc)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
